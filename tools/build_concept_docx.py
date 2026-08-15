# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CONCEPT = ROOT / "docs" / "concept"
SHEET = ROOT / "assets" / "icons" / "sheet.png"
OUT = ROOT / "docs" / "Efir86-koncept.docx"

INK = RGBColor(0x1A, 0x16, 0x12)
RUST = RGBColor(0x8A, 0x3A, 0x28)
MUTED = RGBColor(0x5A, 0x52, 0x48)

FRAMES = [
    ("01-screen.jpg", "Главный кадр. Схема, лента фур, рация. Те же штампы, что в Godot."),
    ("02-call.jpg", "Живой вызов. Дырки в тексте, волна, две кнопки. Частота Б — земля."),
    ("03-stop.jpg", "Остановка. Короткий вдох, три накладные, потом пыль снова идёт."),
    ("04-forks.jpg", "Ночь целиком — три вилки. В демо только первая: карьер → 14."),
    ("05-tower.jpg", "Вышка 14. Связной садится. Демо-срез на этом закрывается."),
    ("06-result.jpg", "Итог — накладная, не проценты. Что проехал, кого услышал."),
]


def set_run_font(run, name="Calibri", size=12, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    set_run_font(run, size=9, color=MUTED)


def p(doc, text, *, size=12, bold=False, italic=False, color=INK, space_after=10, space_before=0, align="left"):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.15
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    set_run_font(run, size=16, bold=True, color=INK)


def picture(doc, path: Path, caption: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    para.add_run().add_picture(str(path), width=Cm(15.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_run_font(r, size=10, italic=True, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Эфир86  ·  концепт")
    set_run_font(hr, size=9, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("черновик  ·  штампы из игры  ·  ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(fp)

    p(doc, "ЭФИР 86", size=28, bold=True, align="center", space_after=4, space_before=8)
    p(doc, "Ночное радио колонны", size=16, italic=True, color=RUST, align="center", space_after=8)
    p(
        doc,
        "Рассказ об игре, чтобы прочитать вслух и посмотреть кадры. Кадры собраны из тех же штампов, что стоят в Godot — сеттинг один, картинки не из другого мира.",
        size=11,
        italic=True,
        color=MUTED,
        align="center",
        space_after=16,
    )

    heading(doc, "В двух словах")
    p(doc, "Ты дежуришь на рации. Рядом колонна фур. Сзади пыль закрывает дорогу — назад нельзя. Впереди посёлки и старые вышки. Звонят свои из кабин и люди с трассы. Рук мало. Иногда голос врёт.", align="justify")
    p(doc, "Не боевик про пустыню и не диспетчер настоящего города. Одна ночь: слышать, решать, кого послать, дотянуть колонну до живой вышки.", align="justify")

    heading(doc, "Мир")
    p(doc, "Обычная связь мертва. Остались пыль с карьеров, столбы с репитерами и люди на частоте. Посёлки живые и злые. Кто-то кормит эфир шумом: тишину боятся сильнее пыли.", align="justify")
    p(doc, "Ты не герой с лицом. Ты смена. Страх простой: не расслышал и послал не туда. Пыль бьёт по времени. Глушение — по словам и по карте.", align="justify")

    heading(doc, "Что делаешь")
    p(doc, "На перегоне слушаешь, ставишь рычаг (свои или земля), жмёшь одну из двух кнопок, шлёшь человека с ленты. Пока он едет — может упасть второй вызов.", align="justify")
    p(doc, "На остановке короткий вдох: три бумажки и две дороги. Потом пыль снова ползёт. Берёшь фуру, договор с местом или вещь на всю ночь.", align="justify")
    p(doc, "Каждая ночь собирается заново. Законы те же, расклад другой. Не доехал — другая ночь, не та же карта наизусть.", align="justify")

    heading(doc, "Как выглядит")
    p(doc, "Один кадр. Сверху схема. Снизу лента коробок. Сбоку железная рация. Восемь красок. Живой только зелёный на волне. Лиц и 3D нет.", align="justify")
    p(doc, "Штампы рисуем сами, одним генератором. Если значок не читается — правим геометрию, не заказываем картину.", align="justify")

    heading(doc, "Кадры")
    for name, cap in FRAMES:
        picture(doc, CONCEPT / name, cap)

    heading(doc, "Все штампы разом")
    if SHEET.exists():
        picture(doc, SHEET, "Фуры, фишки, узлы, рация, вызовы. Один размер, одна палитра.")

    heading(doc, "Демо сейчас")
    p(doc, "В Godot уже открывается срез: карьер → Заправка или Решётка → вышка 14. Берёшь накладную, кликаешь дорогу, ловишь вызов, шлёшь кружок, доезжаешь. На 14 садится связной.", align="justify")
    p(doc, "В срез не вошли: вся ночь, третья частота, бунт своих, выезд с мачты, голоса. Сначала этот кусок должен честно играться.", align="justify")

    heading(doc, "Полная ночь (когда срез встанет)")
    p(doc, "Три вилки: пустая быстрая дорога или густые дворы. Земля помнит, как проехал. Антенна рисует правду на карте — сдохла, карта врёт. Свои могут отказать, если радио среди колонны уже никто не слушает.", align="justify")
    p(doc, "Доехал до мачты — пыль встаёт, слова впервые целые. Можно стоять. Можно глупо уехать назад в пыль за одним голосом.", align="justify")

    heading(doc, "Кто делает")
    p(doc, "Трое. Код, логика, Godot, звук — нормально. Арт — узкое место, поэтому игра живёт как панель, не как фильм. Этот документ и кадры можно показывать другу: вот во что играем, вот как это выглядит без обещания кино.", align="justify")

    p(doc, "Кого спасти сейчас, чтобы потом вообще доехать.", size=14, italic=True, bold=True, align="center", space_before=18)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
